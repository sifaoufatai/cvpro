import json
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def set_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    add_page_number(run)

def add_contact_info(doc, contact):
    # Add name
    name_paragraph = doc.add_paragraph()
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_paragraph.add_run(contact['name'])
    name_run.bold = True
    name_run.font.size = Pt(16)
    
    # Add contact details
    contact_paragraph = doc.add_paragraph()
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = f"{contact['email']} | {contact['phone']} | {contact['linkedin']}"
    if 'location' in contact:
        contact_text += f" | {contact['location']}"
    contact_run = contact_paragraph.add_run(contact_text)
    contact_run.font.size = Pt(9)
    
    # Add a line
    doc.add_paragraph().add_run().add_break()

def add_section_header(doc, title):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)
    # Add a line under the section header
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.border_bottom_color = RGBColor(0, 51, 102)
    paragraph.paragraph_format.border_bottom_width = Pt(1)

def add_summary(doc, summary):
    add_section_header(doc, 'Professional Summary')
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(summary)
    run.font.size = Pt(10)

def add_skills(doc, skills):
    add_section_header(doc, 'Skills')
    
    # Split skills into two columns
    mid = (len(skills) + 1) // 2
    col1 = skills[:mid]
    col2 = skills[mid:]
    
    # Create a table with 2 columns
    table = doc.add_table(rows=max(len(col1), len(col2)), cols=2)
    table.style = 'Table Grid'
    
    # Add skills to the table
    for i in range(len(col1)):
        cell = table.cell(i, 0)
        cell.text = f"• {col1[i]}"
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        
    for i in range(len(col2)):
        cell = table.cell(i, 1)
        cell.text = f"• {col2[i]}"
        cell.paragraphs[0].runs[0].font.size = Pt(10)
    
    # Adjust table style
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
    
    doc.add_paragraph()  # Add some space after the table

def add_experience(doc, experience):
    add_section_header(doc, 'Professional Experience')
    
    for exp in experience:
        # Add job title
        title_paragraph = doc.add_paragraph()
        title_paragraph.paragraph_format.space_after = Pt(0)
        title_run = title_paragraph.add_run(exp['title'])
        title_run.bold = True
        title_run.font.size = Pt(11)
        
        # Add company, location, and dates
        details_paragraph = doc.add_paragraph()
        details_paragraph.paragraph_format.space_after = Pt(0)
        details_text = f"{exp['company']} | {exp['location']} | {exp['dates']}"
        details_run = details_paragraph.add_run(details_text)
        details_run.italic = True
        details_run.font.size = Pt(9)
        
        # Add achievements
        for achievement in exp['achievements']:
            paragraph = doc.add_paragraph(style='List Bullet')
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(achievement)
            run.font.size = Pt(10)
        
        doc.add_paragraph()  # Add space between experiences

def add_education(doc, education):
    add_section_header(doc, 'Education')
    
    for edu in education:
        # Add degree
        degree_paragraph = doc.add_paragraph()
        degree_paragraph.paragraph_format.space_after = Pt(0)
        degree_run = degree_paragraph.add_run(edu['degree'])
        degree_run.bold = True
        degree_run.font.size = Pt(10)
        
        # Add institution and year
        details_paragraph = doc.add_paragraph()
        details_paragraph.paragraph_format.space_after = Pt(6)
        details_text = f"{edu['institution']} | {edu['year']}"
        details_run = details_paragraph.add_run(details_text)
        details_run.italic = True
        details_run.font.size = Pt(9)

def add_projects(doc, projects):
    add_section_header(doc, 'Projects')
    
    for project in projects:
        # Add project name
        name_paragraph = doc.add_paragraph()
        name_paragraph.paragraph_format.space_after = Pt(0)
        name_run = name_paragraph.add_run(project['name'])
        name_run.bold = True
        name_run.font.size = Pt(10)
        
        # Add technologies
        tech_paragraph = doc.add_paragraph()
        tech_paragraph.paragraph_format.space_after = Pt(0)
        tech_run = tech_paragraph.add_run(project['technologies'])
        tech_run.italic = True
        tech_run.font.size = Pt(9)
        
        # Add description
        desc_paragraph = doc.add_paragraph()
        desc_paragraph.paragraph_format.space_after = Pt(6)
        desc_run = desc_paragraph.add_run(project['description'])
        desc_run.font.size = Pt(10)

def add_languages(doc, languages):
    add_section_header(doc, 'Languages')
    
    lang_text = " • ".join([f"{lang['language']} ({lang['level']})" for lang in languages])
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(lang_text)
    run.font.size = Pt(10)

def generate_resume(json_file, output_file):
    # Load resume data
    with open(json_file, 'r', encoding='utf-8') as f:
        resume_data = json.load(f)
    
    # Create a new document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Add footer with page number
    set_footer(doc)
    
    # Add content
    add_contact_info(doc, resume_data['contact_info'])
    add_summary(doc, resume_data['summary'])
    add_skills(doc, resume_data['skills'])
    
    if 'experience' in resume_data and resume_data['experience']:
        add_experience(doc, resume_data['experience'])
    
    if 'education' in resume_data and resume_data['education']:
        add_education(doc, resume_data['education'])
    
    if 'projects' in resume_data and resume_data['projects']:
        add_projects(doc, resume_data['projects'])
    
    if 'languages' in resume_data and resume_data['languages']:
        add_languages(doc, resume_data['languages'])
    
    # Save the document
    doc.save(output_file)
    print(f"Resume generated successfully at: {output_file}")

if __name__ == "__main__":
    # Get the directory of the current script
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    # Define input and output paths
    json_file = os.path.join(script_dir, 'resume_ats_optimized.json')
    output_file = os.path.join(script_dir, 'FATAI_IDRISSOU_Resume.docx')
    
    # Generate the resume
    generate_resume(json_file, output_file)
